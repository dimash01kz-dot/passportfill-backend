from fastapi import FastAPI, File, UploadFile, HTTPException, Header, Body
from fastapi.middleware.cors import CORSMiddleware
import anthropic
import base64
import os
import json
import traceback
import httpx
from typing import Optional

app = FastAPI(title="PassportFill API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY")

PASSPORT_PROMPT = """Ты эксперт по чтению паспортов. Извлеки все данные из этого паспорта и верни ТОЛЬКО JSON без лишнего текста.

Верни JSON в таком формате:
{
  "last_name": "ФАМИЛИЯ латиницей",
  "first_name": "ИМЯ латиницей",
  "last_name_cyr": "Фамилия на кириллице (если есть в паспорте, обычно на казахском/русском)",
  "first_name_cyr": "Имя на кириллице (если есть в паспорте)",
  "middle_name": "ОТЧЕСТВО если есть",
  "birth_date": "ДД.ММ.ГГГГ",
  "passport_series": "серия (обычно одна буква N для казахских)",
  "passport_number": "номер паспорта (только цифры)",
  "expire_date": "ДД.ММ.ГГГГ дата окончания",
  "issue_date": "ДД.ММ.ГГГГ дата выдачи паспорта",
  "issued_by": "кем выдан паспорт латинскими буквами (например MVDKAZ или MIA RK)",
  "iin": "ИИН 12 цифр если есть",
  "gender": "M или F",
  "citizenship": "KAZ или RUS и т.д.",
  "tourist_type": "MR для мужчины взрослого, MRS для женщины взрослой, CHD для ребёнка до 12 лет",
  "nationality": "национальность"
}

Если какого-то поля нет — поставь null. Верни ТОЛЬКО JSON, без объяснений."""


async def amount_to_words_kzt(amount):
    """Convert KZT amount to Russian words"""
    try:
        from num2words import num2words
        amount_int = int(float(str(amount).replace(",", ".").replace(" ", "")))
        words = num2words(amount_int, lang='ru')
        words = words[0].upper() + words[1:]
        return f"{words} тенге"
    except Exception:
        return ""


async def record_history(profile_id: str, count: int, operator: str = None):
    """Record processing in history table"""
    try:
        async with httpx.AsyncClient() as client:
            await client.post(
                f"{SUPABASE_URL}/rest/v1/processing_history",
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    "Content-Type": "application/json"
                },
                json={"profile_id": profile_id, "tourists_count": count, "operator": operator},
                timeout=5.0
            )
    except:
        pass


async def check_and_use_credit(api_key: str) -> dict:
    """Check API key in Supabase and deduct 1 credit"""
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        # If Supabase not configured, allow all requests (dev mode)
        print("WARNING: Supabase not configured, skipping credit check")
        return {"success": True, "credits_left": 999}

    async with httpx.AsyncClient() as client:
        response = await client.post(
            f"{SUPABASE_URL}/rest/v1/rpc/use_credit",
            headers={
                "apikey": SUPABASE_SERVICE_KEY,
                "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                "Content-Type": "application/json"
            },
            json={"p_api_key": api_key},
            timeout=10.0
        )

        if response.status_code != 200:
            print(f"Supabase error: {response.status_code} {response.text}")
            return {"success": False, "error": "Ошибка проверки ключа"}

        return response.json()


@app.get("/")
def root():
    return {"status": "PassportFill API работает", "version": "1.0"}


@app.post("/extract")
async def extract_passport(
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(None)
):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="API ключ не передан")

    api_key = authorization.replace("Bearer ", "").strip()
    if not api_key:
        raise HTTPException(status_code=401, detail="Неверный API ключ")

    # Check credits in Supabase
    credit_result = await check_and_use_credit(api_key)
    if not credit_result.get("success"):
        error = credit_result.get("error", "Ошибка")
        if "кредит" in error.lower() or "недостаточно" in error.lower():
            raise HTTPException(status_code=402, detail="Недостаточно кредитов. Пополните баланс на passportfill.kz")
        raise HTTPException(status_code=401, detail=f"Неверный API ключ: {error}")

    print(f"Processing file: {file.filename}, type: {file.content_type}, credits_left: {credit_result.get('credits_left')}")

    file_bytes = await file.read()

    if len(file_bytes) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Файл слишком большой (макс 10МБ)")

    content_type = file.content_type or "image/jpeg"
    if content_type not in ["image/jpeg", "image/png", "image/webp", "application/pdf"]:
        print(f"Unsupported type: {content_type}, forcing image/jpeg")
        content_type = "image/jpeg"

    file_b64 = base64.standard_b64encode(file_bytes).decode("utf-8")

    try:
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        print("Calling Claude API...")

        if content_type == "application/pdf":
            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=1024,
                messages=[{"role": "user", "content": [
                    {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": file_b64}},
                    {"type": "text", "text": PASSPORT_PROMPT}
                ]}]
            )
        else:
            message = client.messages.create(
                model="claude-opus-4-5",
                max_tokens=1024,
                messages=[{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": content_type, "data": file_b64}},
                    {"type": "text", "text": PASSPORT_PROMPT}
                ]}]
            )

        response_text = message.content[0].text.strip()
        print(f"Claude response: {response_text[:200]}")

        if response_text.startswith("```"):
            response_text = response_text.split("```")[1]
            if response_text.startswith("json"):
                response_text = response_text[4:]

        passport_data = json.loads(response_text.strip())
        del file_bytes
        del file_b64

        return {
            "success": True,
            "data": passport_data,
            "credits_left": credit_result.get("credits_left"),
            "total_processed": credit_result.get("total_processed")
        }

    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        raise HTTPException(status_code=422, detail="Не удалось прочитать паспорт. Попробуйте более чёткое фото.")
    except anthropic.APIError as e:
        print(f"Anthropic API error: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка AI: {str(e)}")
    except Exception as e:
        print(f"Unexpected error: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Ошибка сервера: {str(e)}")


@app.post("/record_history")
async def record_history_endpoint(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None)
):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="API ключ не передан")
    
    api_key = authorization.replace("Bearer ", "").strip()
    count = body.get("count", 1)
    operator = body.get("operator", "Тур оператор")

    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        return {"success": True}

    try:
        # Get profile_id by api_key
        async with httpx.AsyncClient() as client:
            resp = await client.get(
                f"{SUPABASE_URL}/rest/v1/profiles?api_key=eq.{api_key}&select=id",
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}"
                },
                timeout=5.0
            )
            profiles = resp.json()
            if not profiles:
                return {"success": False}
            
            profile_id = profiles[0]["id"]
            
            # Insert history record
            await client.post(
                f"{SUPABASE_URL}/rest/v1/processing_history",
                headers={
                    "apikey": SUPABASE_SERVICE_KEY,
                    "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    "Content-Type": "application/json"
                },
                json={"profile_id": profile_id, "tourists_count": count, "operator": operator},
                timeout=5.0
            )
        return {"success": True}
    except Exception as e:
        print(f"History error: {e}")
        return {"success": False}


@app.get("/health")
def health():
    return {"ok": True}


# ── Google API Setup ──────────────────────────────────────────
import re
from datetime import datetime
from google.oauth2 import service_account
from googleapiclient.discovery import build

GOOGLE_CREDENTIALS_JSON = os.environ.get("GOOGLE_CREDENTIALS_JSON")
SPREADSHEET_ID = "1tbQXU_gQwiY_drbO4cWt0-Jg4gK0PY47wGB_1TEUleU"
TEMPLATE_DOC_ID = "1XSjOVbCFwykLVi5AhRV_HtfDK-zo4XMB"

def get_google_services():
    if not GOOGLE_CREDENTIALS_JSON:
        return None, None
    creds_dict = json.loads(GOOGLE_CREDENTIALS_JSON)
    creds = service_account.Credentials.from_service_account_info(
        creds_dict,
        scopes=[
            "https://www.googleapis.com/auth/documents",
            "https://www.googleapis.com/auth/spreadsheets",
            "https://www.googleapis.com/auth/drive"
        ]
    )
    docs = build("docs", "v1", credentials=creds)
    sheets = build("sheets", "v4", credentials=creds)
    drive = build("drive", "v3", credentials=creds)
    return docs, sheets, drive


@app.post("/create_contract")
async def create_contract(
    body: dict = Body(...),
    authorization: Optional[str] = Header(None)
):
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="API ключ не передан")

    api_key = authorization.replace("Bearer ", "").strip()
    tour_data = body.get("tour", {})
    tourists = body.get("tourists", [])
    contract_num = body.get("contract_num", "")
    agency = body.get("agency", {})

    if not tourists:
        raise HTTPException(status_code=400, detail="Нет данных туристов")

    try:
        docs_svc, sheets_svc, drive_svc = get_google_services()
        if not docs_svc:
            raise HTTPException(status_code=500, detail="Google API не настроен")

        # 1. Download template and fill with python-docx
        from docx import Document
        import googleapiclient.http
        import io

        today = datetime.now().strftime("%d.%m.%Y")

        # Download template as docx (it's already a .docx file, not Google Doc)
        request = drive_svc.files().get_media(fileId=TEMPLATE_DOC_ID)
        file_buffer = io.BytesIO()
        downloader = googleapiclient.http.MediaIoBaseDownload(file_buffer, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        file_buffer.seek(0)

        # 2. Fill placeholders using python-docx
        main_tourist = tourists[0]
        tourist_fio = f"{main_tourist.get('last_name', '')} {main_tourist.get('first_name', '')}".strip()
        tourist_fio_cyr = f"{main_tourist.get('last_name_cyr', '') or main_tourist.get('last_name', '')} {main_tourist.get('first_name_cyr', '') or main_tourist.get('first_name', '')}".strip()
        tourist_iin = main_tourist.get("iin", "")
        copy_title = f"Договор {tourist_fio_cyr} №{contract_num} от {today}".strip()
        price_kzt = tour_data.get("price_kzt", "")
        price_currency = tour_data.get("price_currency", "USD")
        price_currency_val = tour_data.get("price_currency_val", "")
        flights = tour_data.get("flights", [])

        replacements = {
            "{{НОМЕР_ДОГОВОРА}}": str(contract_num),
            "{{СТОРОНА_2}}": f"{tourist_fio_cyr}, ИИН №{tourist_iin}",
            "{{СУММА_ТНГ}}": str(price_kzt),
            "{{СУММА_ПРОПИСЬЮ}}": tour_data.get("price_words") or amount_to_words_kzt(price_kzt),
            "{{СУММА_ВАЛЮТА}}": f"{price_currency_val} {price_currency}",
            "{{ФИО_ТУРИСТА}}": tourist_fio_cyr,
            "{{ДАТА}}": today,
            "{{СТРАНА_ГОРОД}}": f"{tour_data.get('country','')} / {tour_data.get('city','')}",
            "{{ДАТА_НАЧАЛА}}": tour_data.get("date_start", ""),
            "{{ДАТА_ОКОНЧАНИЯ}}": tour_data.get("date_end", ""),
            "{{ОТЕЛЬ}}": tour_data.get("hotel_name", ""),
            "{{ЗВЕЗДЫ}}": tour_data.get("hotel_stars", ""),
            "{{СТРАНА}}": tour_data.get("country", ""),
            "{{ТИП_НОМЕРА}}": tour_data.get("room_type", ""),
            "{{ПИТАНИЕ}}": tour_data.get("meal_type", ""),
            "{{ДАТЫ_ТУРА}}": f"{tour_data.get('date_start','')} - {tour_data.get('date_end','')}",
            "{{КОЛ_ТУРИСТОВ}}": str(len(tourists)),
            "{{СТРАХОВКА}}": tour_data.get("insurance", ""),
            "{{ТРАНСФЕР}}": tour_data.get("transfer", ""),
            "{{МАРШРУТ_1}}": flights[0].get("route", "") if len(flights) > 0 else "",
            "{{АВИАКОМПАНИЯ_1}}": flights[0].get("airline", "") if len(flights) > 0 else "",
            "{{НОМЕР_РЕЙСА_1}}": flights[0].get("number", "") if len(flights) > 0 else "",
            "{{ВРЕМЯ_1}}": flights[0].get("time", "") if len(flights) > 0 else "",
            "{{ДАТА_РЕЙСА_1}}": flights[0].get("date", "") if len(flights) > 0 else "",
            "{{КЛАСС_1}}": flights[0].get("class", "") if len(flights) > 0 else "",
            "{{МАРШРУТ_2}}": flights[1].get("route", "") if len(flights) > 1 else "",
            "{{АВИАКОМПАНИЯ_2}}": flights[1].get("airline", "") if len(flights) > 1 else "",
            "{{НОМЕР_РЕЙСА_2}}": flights[1].get("number", "") if len(flights) > 1 else "",
            "{{ВРЕМЯ_2}}": flights[1].get("time", "") if len(flights) > 1 else "",
            "{{ДАТА_РЕЙСА_2}}": flights[1].get("date", "") if len(flights) > 1 else "",
            "{{КЛАСС_2}}": flights[1].get("class", "") if len(flights) > 1 else "",
        }

        def replace_text(text):
            for k, v in replacements.items():
                text = text.replace(k, str(v) if v else "")
            return text

        doc = Document(file_buffer)

        # First pass: replace simple placeholders everywhere (paragraphs + table cells)
        for para in doc.paragraphs:
            for run in para.runs:
                run.text = replace_text(run.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = replace_text(run.text)

        # Second pass: find passport data table row and fill/duplicate per tourist
        import copy
        from docx.table import _Row

        def fill_passport_row(row, tourist):
            values = [
                f"{tourist.get('last_name','')} {tourist.get('first_name','')}".strip(),
                tourist.get('birth_date', ''),
                f"{tourist.get('passport_series','')}{tourist.get('passport_number','')}".strip(),
                tourist.get('expire_date', ''),
                tourist.get('citizenship', 'KAZ'),
            ]
            for i, cell in enumerate(row.cells):
                if i >= len(values):
                    break
                if cell.paragraphs:
                    p = cell.paragraphs[0]
                    # clear extra paragraphs
                    for extra_p in cell.paragraphs[1:]:
                        extra_p._element.getparent().remove(extra_p._element)
                    if p.runs:
                        p.runs[0].text = values[i]
                        for extra_run in p.runs[1:]:
                            extra_run.text = ""
                    else:
                        p.add_run(values[i])
                else:
                    cell.add_paragraph(values[i])

        for table in doc.tables:
            target_row = None
            for row in table.rows:
                row_text = "".join(cell.text for cell in row.cells)
                if "{{ПАСПОРТНЫЕ_ДАННЫЕ}}" in row_text:
                    target_row = row
                    break
            if target_row is None:
                continue

            # Fill first tourist into the template row
            fill_passport_row(target_row, tourists[0])

            # Duplicate row for remaining tourists
            prev_tr = target_row._tr
            for t in tourists[1:]:
                new_tr = copy.deepcopy(prev_tr)
                prev_tr.addnext(new_tr)
                new_row = _Row(new_tr, table)
                fill_passport_row(new_row, t)
                prev_tr = new_tr
            break

        # 3. Save to buffer
        out_buffer = io.BytesIO()
        doc.save(out_buffer)
        out_buffer.seek(0)
        docx_bytes = out_buffer.read()
        docx_b64 = base64.standard_b64encode(docx_bytes).decode("utf-8")

        # 4. Add row to Google Sheets
        if sheets_svc:
            operator_price_kzt = tour_data.get("operator_price_kzt", "")
            operator_price_currency_val = tour_data.get("operator_price_currency_val", "")
            operator_price_currency = tour_data.get("operator_price_currency", "")

            income = ""
            try:
                if price_kzt and operator_price_kzt:
                    income = str(round(float(str(price_kzt).replace(",", ".")) - float(str(operator_price_kzt).replace(",", "."))))
            except Exception:
                income = ""

            row_data = [
                "", today,
                tour_data.get("date_start", ""), tour_data.get("date_end", ""),
                body.get("trip_type", "пакетный тур"), "",
                tour_data.get("country", ""), tourist_fio_cyr,
                str(len(tourists)), agency.get("manager", ""),
                agency.get("status", "в работе"), tourist_iin,
                tour_data.get("operator", ""), "",
                str(price_kzt), f"{price_currency_val} {price_currency}",
                str(operator_price_kzt), f"{operator_price_currency_val} {operator_price_currency}",
                income,
                f"N{contract_num} от {today}", tour_data.get("hotel_name", ""),
                body.get("note", ""), body.get("tourist_phone", ""),
            ]
            sheets_svc.spreadsheets().values().append(
                spreadsheetId=SPREADSHEET_ID,
                range="2025!A:W",
                valueInputOption="USER_ENTERED",
                body={"values": [row_data]}
            ).execute()

        return {
            "success": True,
            "file_base64": docx_b64,
            "filename": copy_title + ".docx",
            "title": copy_title
        }

    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Ошибка создания договора: {str(e)}")
